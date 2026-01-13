"""Reporting helpers: Word generation, overlays, and run export packaging."""
from __future__ import annotations

import hashlib
import json
import re
import shutil
from collections import defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable
from zipfile import ZIP_DEFLATED, ZipFile

import cv2
from docx import Document
from flask import current_app

from app.extensions import db
from app.models import InferenceRun, Project, Video


class ReportExportError(RuntimeError):
    """Raised when an inference run cannot be exported."""



def generate_video_report(project: Project, video: Video, inference_run: InferenceRun | None = None) -> Path:
    """Generate a minimal Word report summarizing detections for a video."""
    if inference_run is None:
        inference_run = (
            InferenceRun.query.filter_by(video_id=video.id, status="completed")
            .order_by(InferenceRun.created_at.desc())
            .first()
        )

    report_dir = Path("reports") / f"project_{project.id}"
    report_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    report_path = report_dir / f"video_{video.id}_report_{timestamp}.docx"

    doc = Document()
    doc.add_heading(f"Video Logo Detection Report", level=1)

    # Project details
    doc.add_heading("Project", level=2)
    doc.add_paragraph(f"Name: {project.name}")
    doc.add_paragraph(f"Description: {project.description}")
    doc.add_paragraph(f"Budget Limit: {project.budget_limit} {project.currency}")

    # Video details
    doc.add_heading("Video", level=2)
    doc.add_paragraph(f"Video ID: {video.id}")
    doc.add_paragraph(f"Source: {video.original_path}")
    doc.add_paragraph(f"Resolution: {video.resolution or 'Unknown'}")
    doc.add_paragraph(f"Duration (s): {video.duration_seconds or 'Unknown'}")

    if inference_run is None:
        doc.add_paragraph("No inference runs available for this video.")
    else:
        doc.add_heading("Inference Summary", level=2)
        doc.add_paragraph(f"Run ID: {inference_run.id}")
        doc.add_paragraph(f"Status: {inference_run.status}")
        doc.add_paragraph(f"Models: {', '.join(inference_run.model_ids or [])}")

        detections_by_model: dict[str, list] = {}
        for detection in inference_run.detections:
            detections_by_model.setdefault(detection.model_id or "unknown", []).append(detection)

        for model_id, detections in detections_by_model.items():
            doc.add_heading(f"Model: {model_id}", level=3)
            table = doc.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Frame"
            hdr_cells[1].text = "Timestamp (s)"
            hdr_cells[2].text = "Label"
            hdr_cells[3].text = "Confidence"
            for detection in detections:
                row_cells = table.add_row().cells
                row_cells[0].text = str(detection.frame_index or 0)
                row_cells[1].text = f"{float(detection.timestamp_seconds or 0):.2f}"
                row_cells[2].text = detection.label or ""
                row_cells[3].text = f"{float(detection.confidence or 0):.2f}"

    doc.save(report_path)
    db.session.commit()
    return report_path


def _hex_to_bgr(color: str) -> tuple[int, int, int]:
    if not isinstance(color, str) or not color.startswith("#") or len(color) != 7:
        return (0, 0, 0)
    r = int(color[1:3], 16)
    g = int(color[3:5], 16)
    b = int(color[5:7], 16)
    return (b, g, r)


def _model_color_bgr(model_id: str) -> tuple[int, int, int]:
    mapping = {
        "A": _hex_to_bgr("#FF0000"),
        "B": _hex_to_bgr("#003366"),
    }
    return mapping.get((model_id or "").upper(), (0, 255, 255))


def _clamp01(value: float | int | None) -> float:
    try:
        v = float(value or 0)
    except (TypeError, ValueError):
        return 0.0
    return max(0.0, min(1.0, v))


def _bbox_pixels(bbox: dict, width: int, height: int) -> tuple[int, int, int, int]:
    left = _clamp01(bbox.get("left"))
    top = _clamp01(bbox.get("top"))
    right = _clamp01(bbox.get("right"))
    bottom = _clamp01(bbox.get("bottom"))
    x1 = int(left * max(width, 1))
    y1 = int(top * max(height, 1))
    x2 = int(right * max(width, 1))
    y2 = int(bottom * max(height, 1))
    return x1, y1, x2, y2


def draw_frame_overlay(
    image_path: str | Path,
    detections: list[dict],
    output_path: str | Path,
    thickness: int = 2,
    *,
    return_size: bool = False,
) -> Path | tuple[Path, tuple[int, int]]:
    """Draw bounding boxes for detections onto a frame image and save PNG.

    Colors:
      - Model A: red (#FF0000)
      - Model B: dark blue (#003366)

    Expects detection entries with keys: 'bbox' (normalized left/top/right/bottom),
    'label', 'confidence', and 'model_id' ('A' or 'B').
    """
    img_path = Path(image_path)
    out_path = Path(output_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    image = cv2.imread(str(img_path))
    if image is None:
        raise FileNotFoundError(f"Could not read image at {img_path}")
    h, w = image.shape[:2]

    for det in detections or []:
        bbox = det.get("bbox") or {}
        x1, y1, x2, y2 = _bbox_pixels(bbox, w, h)
        color = _model_color_bgr(str(det.get("model_id", "")))
        cv2.rectangle(image, (x1, y1), (x2, y2), color, thickness)

        label = str(det.get("label") or "")
        conf = det.get("confidence")
        try:
            conf_txt = f"{float(conf):.2f}"
        except (TypeError, ValueError):
            conf_txt = ""
        text = f"{label} ({conf_txt})".strip()
        if text:
            (tw, th), baseline = cv2.getTextSize(text, cv2.FONT_HERSHEY_SIMPLEX, 0.5, 1)
            tx = max(x1, 0)
            ty = max(y1 - 6, th + 4)
            bg_x2 = min(tx + tw + 6, w)
            bg_y2 = ty + th + 4
            roi = image[ty - th - 4 : bg_y2, tx : bg_x2]
            if roi.size > 0:
                overlay = roi.copy()
                cv2.rectangle(overlay, (0, 0), (overlay.shape[1] - 1, overlay.shape[0] - 1), color, -1)
                cv2.addWeighted(overlay, 0.3, roi, 0.7, 0, roi)
                image[ty - th - 4 : bg_y2, tx : bg_x2] = roi
            cv2.putText(image, text, (tx + 3, ty), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (255, 255, 255), 1, cv2.LINE_AA)

    cv2.imwrite(str(out_path), image)
    if return_size:
        return out_path, (w, h)
    return out_path


def _reports_root() -> Path:
    base = current_app.config.get("REPORTS_ROOT", "reports")
    root = Path(base)
    if not root.is_absolute():
        root = Path(current_app.root_path).parent / root
    root.mkdir(parents=True, exist_ok=True)
    return root


def _prepare_run_dir(run_id: int, regenerate: bool) -> Path:
    run_dir = _reports_root() / f"run_{run_id}"
    if regenerate and run_dir.exists():
        shutil.rmtree(run_dir)
    run_dir.mkdir(parents=True, exist_ok=True)
    return run_dir


def _frame_token(frame_index: int) -> str:
    return f"{max(0, int(frame_index)):06d}"


def _sanitize_model_slug(model_id: str | None) -> str:
    value = (model_id or "unknown").strip().lower()
    value = re.sub(r"[^a-z0-9_-]+", "-", value)
    return value.strip("-") or "model"


def _hash_file(path: Path) -> str:
    digest = hashlib.sha256()
    with open(path, "rb") as handle:
        for chunk in iter(lambda: handle.read(65536), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _safe_float(value) -> float | None:
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _clamp_unit(value) -> float:
    numeric = _safe_float(value)
    if numeric is None:
        return 0.0
    return max(0.0, min(1.0, numeric))


def _timestamp_ms(value) -> int:
    numeric = _safe_float(value) or 0.0
    return int(round(max(0.0, numeric) * 1000))


def _normalize_bbox_dict(
    bbox: dict | None,
    *,
    width: int | float | None = None,
    height: int | float | None = None,
) -> dict[str, float]:
    if not isinstance(bbox, dict):
        return {"left": 0.0, "top": 0.0, "right": 0.0, "bottom": 0.0}

    def _maybe_number(value):
        numeric = _safe_float(value)
        return numeric

    left = _maybe_number(bbox.get("left"))
    top = _maybe_number(bbox.get("top"))
    right = _maybe_number(bbox.get("right"))
    bottom = _maybe_number(bbox.get("bottom"))
    if None in (left, top, right, bottom):
        x = _maybe_number(bbox.get("x") or bbox.get("xmin"))
        y = _maybe_number(bbox.get("y") or bbox.get("ymin"))
        box_w = _maybe_number(bbox.get("width"))
        box_h = _maybe_number(bbox.get("height"))
        if None not in (x, y, box_w, box_h):
            left = x
            top = y
            right = x + box_w
            bottom = y + box_h

    frame_w = _safe_float(width) or None
    frame_h = _safe_float(height) or None

    def _normalize(value, scale):
        if value is None:
            return 0.0
        numeric = float(value)
        if scale and scale > 1 and abs(numeric) > 1.0:
            numeric = numeric / scale
        return max(0.0, min(1.0, numeric))

    norm_left = _normalize(left, frame_w)
    norm_top = _normalize(top, frame_h)
    norm_right = _normalize(right, frame_w)
    norm_bottom = _normalize(bottom, frame_h)

    norm_right = max(norm_left, norm_right)
    norm_bottom = max(norm_top, norm_bottom)

    return {
        "left": norm_left,
        "top": norm_top,
        "right": norm_right,
        "bottom": norm_bottom,
    }


def _bbox_pixels_from_normalized(bbox: dict[str, float], width: int, height: int) -> dict[str, int]:
    width = max(int(width), 1)
    height = max(int(height), 1)
    left = bbox.get("left", 0.0)
    top = bbox.get("top", 0.0)
    right = max(left, bbox.get("right", left))
    bottom = max(top, bbox.get("bottom", top))
    x = int(round(left * width))
    y = int(round(top * height))
    box_w = max(1, int(round((right - left) * width)))
    box_h = max(1, int(round((bottom - top) * height)))
    return {"x": x, "y": y, "width": box_w, "height": box_h}


def _model_alias_map(model_ids: Iterable[str] | None) -> dict[str, str]:
    aliases: dict[str, str] = {}
    if not model_ids:
        return aliases
    ordered = list(model_ids)
    if ordered:
        aliases[ordered[0]] = "A"
    if len(ordered) > 1:
        aliases[ordered[1]] = "B"
    return aliases


def _frame_metadata_map(inference_run: InferenceRun) -> dict[int, dict]:
    frames = (inference_run.results or {}).get("frames") or []
    mapping: dict[int, dict] = {}
    for frame in frames:
        index = frame.get("index")
        if index is None:
            continue
        mapping[int(index)] = frame
    return mapping


def _frame_image_path(run: InferenceRun, frame_index: int, metadata: dict) -> Path | None:
    if metadata.get("image_path"):
        return Path(metadata["image_path"])
    video = run.video
    if video is None:
        return None
    media_root = Path(current_app.config.get("PROJECT_MEDIA_ROOT", "media"))
    if not media_root.is_absolute():
        media_root = Path(current_app.root_path).parent / media_root
    return (
        media_root
        / f"project_{video.project_id}"
        / f"video_{video.id}"
        / "runs"
        / f"run_{run.id}"
        / "frames"
        / f"frame_{frame_index:04d}.jpg"
    )


def _write_json(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def _detection_payload(detection, width: int, height: int) -> dict:
    bbox_norm = _normalize_bbox_dict(detection.bbox or {}, width=width, height=height)
    return {
        "label": detection.label or "",
        "confidence": _safe_float(detection.confidence) or 0.0,
        "frameIndex": detection.frame_index,
        "timestampMs": _timestamp_ms(detection.timestamp_seconds),
        "bbox": _bbox_pixels_from_normalized(bbox_norm, width, height),
        "bboxNormalized": bbox_norm,
    }


def build_run_export(
    inference_run: InferenceRun,
    *,
    frame_indices: list[int] | None = None,
    regenerate: bool = True,
) -> Path:
    """Package frame overlays, JSON payloads, and manifest for an inference run."""
    if inference_run is None:
        raise ReportExportError("Inference run is required")
    if inference_run.video is None:
        raise ReportExportError("Inference run is missing its video reference")

    run_dir = _prepare_run_dir(inference_run.id, regenerate)
    frames_dir = run_dir / "frames"
    json_dir = run_dir / "json"
    frames_dir.mkdir(parents=True, exist_ok=True)
    json_dir.mkdir(parents=True, exist_ok=True)

    frame_meta = _frame_metadata_map(inference_run)
    detections = list(inference_run.detections or [])
    if not frame_meta:
        raise ReportExportError("Run has no stored frame metadata to export")
    if not detections:
        raise ReportExportError("Run has no detections to export")

    alias_map = _model_alias_map(inference_run.model_ids)
    vlm_meta = (inference_run.results or {}).get("vlm_overlays") or {}
    vlm_model_id = vlm_meta.get("model_id")
    if vlm_model_id and vlm_model_id not in alias_map:
        alias_map[vlm_model_id] = "B"
    grouped: dict[int, dict[str, list]] = defaultdict(lambda: defaultdict(list))
    for detection in detections:
        if detection.frame_index is None:
            continue
        grouped[int(detection.frame_index)][detection.model_id or "unknown"].append(detection)

    selected_frames = frame_indices or sorted(grouped.keys() or frame_meta.keys())
    if not selected_frames:
        raise ReportExportError("No frames available for export")

    aggregate_payload = {
        "runId": inference_run.id,
        "projectId": inference_run.project_id,
        "videoId": inference_run.video_id,
        "models": inference_run.model_ids or [],
        "frames": [],
    }

    frame_items: list[dict] = []
    json_items: list[dict] = []

    for frame_index in selected_frames:
        metadata = frame_meta.get(frame_index)
        if metadata is None:
            continue
        image_path = _frame_image_path(inference_run, frame_index, metadata)
        if image_path is None or not image_path.is_file():
            raise ReportExportError(f"Frame image missing for frame {frame_index}")

        base_image = cv2.imread(str(image_path))
        if base_image is None:
            raise ReportExportError(f"Unable to read frame image for {frame_index}")
        frame_height, frame_width = base_image.shape[:2]

        overlay_name = f"frame_{_frame_token(frame_index)}_overlay.png"
        overlay_path = frames_dir / overlay_name
        overlay_detections = []
        for model_id, model_detections in grouped.get(frame_index, {}).items():
            alias = alias_map.get(model_id)
            for det in model_detections:
                bbox_norm = _normalize_bbox_dict(det.bbox or {}, width=frame_width, height=frame_height)
                overlay_detections.append(
                    {
                        "bbox": bbox_norm,
                        "label": det.label,
                        "confidence": _safe_float(det.confidence) or 0.0,
                        "model_id": alias or model_id or "unknown",
                    }
                )

        overlay_return = draw_frame_overlay(image_path, overlay_detections, overlay_path, return_size=True)
        if isinstance(overlay_return, tuple):
            _, (reported_width, reported_height) = overlay_return
            if reported_width and reported_height:
                frame_width, frame_height = reported_width, reported_height
        width, height = frame_width, frame_height

        timestamp_seconds = metadata.get("timestamp_seconds")
        timestamp_ms = _timestamp_ms(timestamp_seconds)
        frame_entry = {
            "frameId": frame_index,
            "timestampMs": timestamp_ms,
            "image": f"frames/{overlay_name}",
            "detections": [],
        }

        for model_id, model_detections in grouped.get(frame_index, {}).items():
            alias = alias_map.get(model_id)
            slug = _sanitize_model_slug(alias or model_id)
            json_filename = f"model_{slug}_frame_{_frame_token(frame_index)}.json"
            json_path = json_dir / json_filename
            detection_payloads = [_detection_payload(det, width, height) for det in model_detections]
            payload = {
                "runId": inference_run.id,
                "frameId": frame_index,
                "timestampMs": timestamp_ms,
                "source": {
                    "projectId": inference_run.project_id,
                    "videoId": inference_run.video_id,
                },
                "model": model_id,
                "alias": alias,
                "detections": detection_payloads,
                "diagnostics": {
                    "modelIndex": (inference_run.model_ids or []).index(model_id)
                    if model_id in (inference_run.model_ids or [])
                    else None,
                },
            }
            _write_json(json_path, payload)
            json_items.append({"path": f"json/{json_filename}", "sha256": _hash_file(json_path)})

            for det_payload in detection_payloads:
                frame_entry["detections"].append({
                    "model": model_id,
                    "alias": alias,
                    **det_payload,
                })

        aggregate_payload["frames"].append(frame_entry)
        frame_items.append({"path": f"frames/{overlay_name}", "sha256": _hash_file(overlay_path)})

    if not frame_items:
        raise ReportExportError("No frames could be exported for this run")

    aggregate_path = json_dir / "detections_aggregate.json"
    _write_json(aggregate_path, aggregate_payload)
    json_items.append({"path": "json/detections_aggregate.json", "sha256": _hash_file(aggregate_path)})

    manifest_path = run_dir / "manifest.json"
    manifest = {
        "runId": inference_run.id,
        "projectId": inference_run.project_id,
        "videoId": inference_run.video_id,
        "createdAt": datetime.now(timezone.utc).isoformat(),
        "models": inference_run.model_ids or [],
        "clip": (inference_run.results or {}).get("clip"),
        "counts": {"frames": len(frame_items), "json": len(json_items)},
        "items": {
            "frames": frame_items,
            "json": json_items,
        },
    }
    _write_json(manifest_path, manifest)

    archive_path = _reports_root() / f"run_{inference_run.id}.zip"
    if archive_path.exists():
        archive_path.unlink()
    with ZipFile(archive_path, "w", ZIP_DEFLATED) as archive:
        for path in run_dir.rglob("*"):
            if path.is_file():
                arcname = Path(run_dir.name) / path.relative_to(run_dir)
                archive.write(path, arcname=str(arcname).replace("\\", "/"))

    return archive_path
