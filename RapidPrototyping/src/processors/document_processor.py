"""
Document Processor

Handles processing of various document formats for analysis.
"""

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from src.clients.clarifai_client import ClarifaiClient
from src.config import get_config


logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """A chunk of processed document."""
    content: str
    source: str
    page: Optional[int] = None
    section: Optional[str] = None
    metadata: Dict[str, Any] = None


@dataclass 
class ProcessedDocument:
    """Result of document processing."""
    filename: str
    file_type: str
    chunks: List[DocumentChunk]
    full_text: str
    metadata: Dict[str, Any]
    summary: Optional[str] = None


class DocumentProcessor:
    """
    Processes various document formats for AI analysis.
    
    Supports:
    - Markdown (.md)
    - Plain text (.txt)
    - PDF (.pdf)
    - Word documents (.docx)
    """
    
    def __init__(self, client: Optional[ClarifaiClient] = None):
        """
        Initialize the document processor.
        
        Args:
            client: Optional ClarifaiClient for AI-powered processing.
        """
        self.client = client
        self.config = get_config()
    
    def process(
        self,
        file_path: Union[str, Path],
        extract_sections: bool = True,
        generate_summary: bool = False,
        chunk_size: int = 2000,
    ) -> ProcessedDocument:
        """
        Process a document file.
        
        Args:
            file_path: Path to the document.
            extract_sections: Whether to extract document sections.
            generate_summary: Whether to generate AI summary.
            chunk_size: Maximum characters per chunk.
            
        Returns:
            ProcessedDocument with extracted content.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        suffix = path.suffix.lower()
        
        # Extract text based on file type
        if suffix == '.md':
            full_text, metadata = self._process_markdown(path)
        elif suffix == '.txt':
            full_text, metadata = self._process_text(path)
        elif suffix == '.pdf':
            full_text, metadata = self._process_pdf(path)
        elif suffix == '.docx':
            full_text, metadata = self._process_docx(path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
        
        # Create chunks
        chunks = self._create_chunks(full_text, str(path), chunk_size)
        
        # Extract sections if requested
        if extract_sections:
            chunks = self._extract_sections(chunks, suffix)
        
        # Generate summary if requested
        summary = None
        if generate_summary and self.client:
            summary = self._generate_summary(full_text)
        
        return ProcessedDocument(
            filename=path.name,
            file_type=suffix,
            chunks=chunks,
            full_text=full_text,
            metadata=metadata,
            summary=summary,
        )
    
    def _process_markdown(self, path: Path) -> tuple[str, Dict[str, Any]]:
        """Process a Markdown file."""
        try:
            import frontmatter
            post = frontmatter.load(path)
            return post.content, dict(post.metadata)
        except ImportError:
            # Fallback without frontmatter parsing
            content = path.read_text(encoding='utf-8')
            return content, {}
    
    def _process_text(self, path: Path) -> tuple[str, Dict[str, Any]]:
        """Process a plain text file."""
        content = path.read_text(encoding='utf-8')
        return content, {"encoding": "utf-8"}
    
    def _process_pdf(self, path: Path) -> tuple[str, Dict[str, Any]]:
        """Process a PDF file."""
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(path)
            pages = []
            for page in reader.pages:
                pages.append(page.extract_text())
            
            metadata = {
                "pages": len(reader.pages),
                "info": reader.metadata if reader.metadata else {},
            }
            
            return "\n\n".join(pages), metadata
        except ImportError:
            logger.warning("PyPDF2 not available. Install with: pip install PyPDF2")
            return "", {"error": "PyPDF2 not installed"}
    
    def _process_docx(self, path: Path) -> tuple[str, Dict[str, Any]]:
        """Process a Word document."""
        try:
            from docx import Document
            
            doc = Document(path)
            paragraphs = [para.text for para in doc.paragraphs]
            
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "sections": len(doc.sections),
            }
            
            return "\n\n".join(paragraphs), metadata
        except ImportError:
            logger.warning("python-docx not available. Install with: pip install python-docx")
            return "", {"error": "python-docx not installed"}
    
    def _create_chunks(
        self,
        text: str,
        source: str,
        chunk_size: int
    ) -> List[DocumentChunk]:
        """Split text into chunks."""
        chunks = []
        
        # Split by paragraphs first
        paragraphs = text.split('\n\n')
        current_chunk = ""
        
        for para in paragraphs:
            if len(current_chunk) + len(para) < chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(DocumentChunk(
                        content=current_chunk.strip(),
                        source=source,
                    ))
                current_chunk = para + "\n\n"
        
        # Don't forget the last chunk
        if current_chunk.strip():
            chunks.append(DocumentChunk(
                content=current_chunk.strip(),
                source=source,
            ))
        
        return chunks
    
    def _extract_sections(
        self,
        chunks: List[DocumentChunk],
        file_type: str
    ) -> List[DocumentChunk]:
        """Extract section information from chunks."""
        if file_type == '.md':
            # Parse Markdown headers
            current_section = None
            for chunk in chunks:
                lines = chunk.content.split('\n')
                for line in lines:
                    if line.startswith('#'):
                        current_section = line.lstrip('#').strip()
                        break
                chunk.section = current_section
        
        return chunks
    
    def _generate_summary(self, text: str, max_length: int = 500) -> str:
        """Generate an AI summary of the document."""
        if not self.client:
            return ""
        
        # Truncate if too long
        if len(text) > 10000:
            text = text[:10000] + "..."
        
        prompt = f"""Summarize the following document concisely, highlighting:
1. Main topics and themes
2. Key requirements or goals mentioned
3. Important technical details
4. Any constraints or considerations

Document:
{text}

Summary:"""
        
        response = self.client.complete(prompt)
        return response.content
    
    def process_multiple(
        self,
        file_paths: List[Union[str, Path]],
        **kwargs
    ) -> List[ProcessedDocument]:
        """
        Process multiple documents.
        
        Args:
            file_paths: List of paths to documents.
            **kwargs: Arguments passed to process().
            
        Returns:
            List of ProcessedDocument objects.
        """
        results = []
        for path in file_paths:
            try:
                doc = self.process(path, **kwargs)
                results.append(doc)
            except Exception as e:
                logger.error(f"Error processing {path}: {e}")
        
        return results
    
    def combine_documents(
        self,
        documents: List[ProcessedDocument],
        generate_combined_summary: bool = False,
    ) -> str:
        """
        Combine multiple documents into a single context.
        
        Args:
            documents: List of processed documents.
            generate_combined_summary: Whether to summarize the combination.
            
        Returns:
            Combined document text.
        """
        parts = []
        
        for doc in documents:
            parts.append(f"## Document: {doc.filename}\n")
            if doc.summary:
                parts.append(f"**Summary:** {doc.summary}\n")
            parts.append(doc.full_text)
            parts.append("\n---\n")
        
        combined = "\n".join(parts)
        
        if generate_combined_summary and self.client:
            summary = self._generate_summary(combined)
            return f"# Combined Document Summary\n\n{summary}\n\n---\n\n{combined}"
        
        return combined
